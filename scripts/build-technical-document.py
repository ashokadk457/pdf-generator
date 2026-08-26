from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "output" / "docs" / "WannaTalk_Intake_PDF_Technical_Guide.docx"
OUT.parent.mkdir(parents=True, exist_ok=True)

NAVY = "17384A"
BLUE = "2E74B5"
GREEN = "19A64A"
MUTED = "607985"
PALE = "E8F1F5"
LIGHT = "F4F7F8"
RED = "9B1C1C"

doc = Document()
sec = doc.sections[0]
sec.page_width = Inches(8.5)
sec.page_height = Inches(11)
sec.top_margin = Inches(0.8)
sec.bottom_margin = Inches(0.75)
sec.left_margin = Inches(0.85)
sec.right_margin = Inches(0.85)
sec.header_distance = Inches(0.35)
sec.footer_distance = Inches(0.35)

styles = doc.styles
normal = styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(10.5)
normal.font.color.rgb = RGBColor.from_string(NAVY)
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.18

for name, size, color, before, after in [
    ("Heading 1", 16, BLUE, 16, 8),
    ("Heading 2", 13, BLUE, 12, 6),
    ("Heading 3", 11.5, NAVY, 9, 4),
]:
    s = styles[name]
    s.font.name = "Calibri"
    s.font.size = Pt(size)
    s.font.bold = True
    s.font.color.rgb = RGBColor.from_string(color)
    s.paragraph_format.space_before = Pt(before)
    s.paragraph_format.space_after = Pt(after)
    s.paragraph_format.keep_with_next = True

for name in ["List Bullet", "List Number"]:
    s = styles[name]
    s.font.name = "Calibri"
    s.font.size = Pt(10.5)
    s.paragraph_format.left_indent = Inches(0.38)
    s.paragraph_format.first_line_indent = Inches(-0.19)
    s.paragraph_format.space_after = Pt(4)
    s.paragraph_format.line_spacing = 1.18

def font(run, size=None, bold=None, color=None, italic=None, name="Calibri"):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    if size: run.font.size = Pt(size)
    if bold is not None: run.bold = bold
    if italic is not None: run.italic = italic
    if color: run.font.color.rgb = RGBColor.from_string(color)
    return run

def shade(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = tcPr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tcPr.append(shd)
    shd.set(qn("w:fill"), fill)

def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = tcPr.first_child_found_in("w:tcMar")
    if tcMar is None:
        tcMar = OxmlElement("w:tcMar")
        tcPr.append(tcMar)
    for m, v in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        node = tcMar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tcMar.append(node)
        node.set(qn("w:w"), str(v)); node.set(qn("w:type"), "dxa")

def set_repeat_table_header(row):
    trPr = row._tr.get_or_add_trPr()
    tblHeader = OxmlElement("w:tblHeader")
    tblHeader.set(qn("w:val"), "true")
    trPr.append(tblHeader)

def set_col_width(cell, width_twips):
    tcPr = cell._tc.get_or_add_tcPr()
    tcW = tcPr.find(qn("w:tcW"))
    tcW.set(qn("w:w"), str(width_twips)); tcW.set(qn("w:type"), "dxa")

def table(headers, rows, widths):
    t = doc.add_table(rows=1, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    t.autofit = False
    t.style = "Table Grid"
    set_repeat_table_header(t.rows[0])
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]; set_col_width(c, widths[i]); shade(c, PALE); set_cell_margins(c)
        c.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = c.paragraphs[0]; p.paragraph_format.space_after = Pt(0)
        font(p.add_run(h), 9.5, True, NAVY)
    for row in rows:
        cells = t.add_row().cells
        for i, value in enumerate(row):
            c = cells[i]; set_col_width(c, widths[i]); set_cell_margins(c)
            p = c.paragraphs[0]; p.paragraph_format.space_after = Pt(0)
            font(p.add_run(str(value)), 9.2, False, NAVY)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    return t

def callout(label, text, fill=LIGHT, accent=GREEN):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.12)
    p.paragraph_format.right_indent = Inches(0.12)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.keep_together = True
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd"); shd.set(qn("w:fill"), fill); pPr.append(shd)
    pBdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left"); left.set(qn("w:val"), "single"); left.set(qn("w:sz"), "18"); left.set(qn("w:color"), accent); left.set(qn("w:space"), "8")
    pBdr.append(left); pPr.append(pBdr)
    font(p.add_run(label + "  "), 10, True, accent)
    font(p.add_run(text), 10, False, NAVY)

def code(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.18)
    p.paragraph_format.right_indent = Inches(0.18)
    p.paragraph_format.space_before = Pt(4); p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.05
    pPr = p._p.get_or_add_pPr(); shd = OxmlElement("w:shd"); shd.set(qn("w:fill"), "F2F4F7"); pPr.append(shd)
    for idx, line in enumerate(text.strip().splitlines()):
        if idx: p.add_run().add_break()
        font(p.add_run(line), 8.5, False, NAVY, name="Consolas")

def bullet(text):
    p = doc.add_paragraph(style="List Bullet"); p.paragraph_format.keep_together = True
    font(p.add_run(text), 10.5, False, NAVY)

def numbered(text):
    p = doc.add_paragraph(style="List Number"); p.paragraph_format.keep_together = True
    font(p.add_run(text), 10.5, False, NAVY)

def page_break():
    doc.add_page_break()

# Running header/footer
h = sec.header.paragraphs[0]
h.alignment = WD_ALIGN_PARAGRAPH.LEFT
font(h.add_run("WannaTalk  |  Developer Technical Guide"), 8.5, True, MUTED)
f = sec.footer.paragraphs[0]
f.alignment = WD_ALIGN_PARAGRAPH.CENTER
font(f.add_run("Confidential - implementation reference"), 8, False, MUTED)

# Cover
p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(30); p.paragraph_format.space_after = Pt(8)
font(p.add_run("WANNATALK"), 11, True, GREEN)
p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(8)
font(p.add_run("Intake Transcript to Clinical PDF"), 26, True, NAVY)
p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(22)
font(p.add_run("Developer Technical Guide"), 16, False, BLUE)
callout("Purpose", "Explain exactly where the OpenAI API is called, how transcript data becomes structured clinical content, and how the final branded PDF is produced and returned to the user.")
table(["Document", "Value"], [
    ("Project", "wannatalk-intake-pdf"),
    ("Runtime", "Node.js, ES modules"),
    ("Primary endpoint", "POST /api/reports"),
    ("Input", "Patient metadata plus UTF-8 .txt transcript"),
    ("Output", "Downloaded PDF and server-side copy under output/pdf/"),
    ("API", "OpenAI Responses API with strict JSON Schema"),
    ("Version", "1.0 - 26 August 2026"),
], [2500, 6860])

doc.add_heading("Document scope", level=1)
for x in [
    "System architecture and end-to-end request flow",
    "Exact OpenAI call location, request shape, and response handling",
    "Extraction rules for supported, inferred, missing, and unassessed data",
    "PDF layout logic, configuration, security, testing, deployment, and troubleshooting",
]: bullet(x)

page_break()
doc.add_heading("1. System overview", level=1)
doc.add_paragraph("The application is a server-rendered workflow with a static browser form. The browser never receives the OpenAI key and never calls OpenAI directly. Express accepts the upload, reads the transcript, calls the extraction module, passes the structured result to the PDF module, and downloads the finished report.")

callout("Core design rule", "Administrative metadata is authoritative from the submitted form. Clinical narrative content is extracted from the transcript. The model is instructed not to invent missing facts, diagnoses, questions, direct observations, or safety assessments.")

doc.add_heading("1.1 End-to-end sequence", level=2)
for step in [
    "User enters patient name, reference number, email and/or phone, and intake type.",
    "User selects a UTF-8 .txt transcript and submits the multipart form.",
    "Express and Multer validate metadata, file type, and the 2 MB size limit.",
    "The server reads the temporary transcript and calls extractReport(transcript).",
    "extractReport sends the transcript to the OpenAI Responses API with a strict JSON Schema.",
    "The validated JSON report is passed to createReportPdf(...).",
    "PDFKit builds the report, applies the optional logo, and saves the PDF under output/pdf/.",
    "Express returns the PDF as a browser download and removes the temporary transcript.",
]: numbered(step)

doc.add_heading("1.2 Component map", level=2)
table(["File", "Responsibility"], [
    ("public/index.html", "Browser form for metadata and transcript upload."),
    ("public/styles.css", "Responsive form styling."),
    ("src/server.js", "HTTP endpoint, validation, upload lifecycle, orchestration, download response."),
    ("src/extract-report.js", "Server-side OpenAI client and structured extraction request."),
    ("src/report-schema.js", "Strict JSON Schema for the clinical report payload."),
    ("src/pdf.js", "Branded PDF layout, pagination, headers, footers, and field formatting."),
    ("scripts/generate-sample.js", "Offline PDF-generation test without an OpenAI call."),
    (".env", "Runtime secrets and deployment-specific configuration; excluded from Git."),
], [2500, 6860])

page_break()
doc.add_heading("2. HTTP API contract", level=1)
doc.add_heading("2.1 Endpoint", level=2)
code("POST /api/reports\nContent-Type: multipart/form-data\nResponse: application/pdf on success")

doc.add_heading("2.2 Request fields", level=2)
table(["Field", "Required", "Source", "Validation"], [
    ("patientName", "Yes", "Form", "Non-empty string"),
    ("referenceNumber", "Yes", "Form", "Non-empty; sanitized for filename"),
    ("email", "Conditional", "Form", "Email or phone must be provided"),
    ("phone", "Conditional", "Form", "Email or phone must be provided"),
    ("intakeType", "Yes", "Form", "text, voice, or mixed"),
    ("transcript", "Yes", "File", ".txt/text/plain, maximum 2 MB, non-empty"),
], [1900, 1300, 1500, 4660])

doc.add_heading("2.3 Example request", level=2)
code('curl -X POST http://localhost:3000/api/reports \\\n+  -F "patientName=Example Patient" \\\n+  -F "referenceNumber=WT-1001" \\\n+  -F "email=patient@example.com" \\\n+  -F "intakeType=text" \\\n+  -F "transcript=@C:/data/intake.txt" \\\n+  --output WannaTalk-WT-1001.pdf')

doc.add_heading("2.4 Responses", level=2)
table(["Status", "Meaning", "Body"], [
    ("200", "Report created", "PDF download"),
    ("400", "Missing metadata, missing/empty transcript, or upload validation failure", "JSON error"),
    ("500", "OpenAI, parsing, filesystem, or PDF-generation failure", "JSON error with details"),
], [1200, 5000, 3160])

page_break()
doc.add_heading("3. Where the OpenAI API is called", level=1)
callout("API call location", "src/extract-report.js, inside extractReport(transcript). This is the only module that creates the OpenAI client or sends transcript content to OpenAI.")

doc.add_heading("3.1 Client initialization", level=2)
code('const client = new OpenAI({\n  apiKey: process.env.OPENAI_API_KEY\n});')
doc.add_paragraph("The SDK reads the secret from the server process. The .env file is excluded by .gitignore, and the browser receives neither the key nor an OpenAI endpoint.")

doc.add_heading("3.2 Responses API request", level=2)
code('const response = await client.responses.create({\n  model: process.env.OPENAI_MODEL || "gpt-5-mini",\n  store: false,\n  instructions: EXTRACTION_RULES,\n  input: transcript,\n  text: {\n    format: {\n      type: "json_schema",\n      name: "laos_intake_report",\n      strict: true,\n      schema: reportSchema\n    }\n  }\n});')

doc.add_heading("3.3 Important request choices", level=2)
for x in [
    "store: false reduces application-state retention for the generated response.",
    "strict: true requires output matching the supplied JSON Schema.",
    "The full transcript is supplied as input; metadata is not sent because it already comes from the trusted form.",
    "The model is configurable through OPENAI_MODEL without changing source code.",
    "response.output_text is parsed with JSON.parse and passed to the PDF layer.",
]: bullet(x)

doc.add_heading("3.4 Extraction instruction policy", level=2)
table(["Situation", "Required behavior"], [
    ("Age/gender explicitly stated", "Extract the value."),
    ("Demographic absent", "Return null; PDF prints Not available."),
    ("Language apparent from writing", "Return inferred and label it as inferred."),
    ("Mood/appearance/behaviour", "Do not claim direct observation from transcript text."),
    ("Safety not directly assessed", "currentSafetyAssessed=false and recommend direct follow-up."),
    ("Questions not recorded", "Return an empty list; PDF says no questions were recorded."),
    ("Evidence", "Use short verbatim transcript excerpts only."),
    ("Diagnosis", "Do not create one; report only diagnoses explicitly stated by the patient."),
], [2650, 6710])

page_break()
doc.add_heading("4. Structured report schema", level=1)
doc.add_paragraph("src/report-schema.js is the contract between OpenAI extraction and PDF generation. additionalProperties is false, required fields are enumerated, and nested risk/evidence/question objects have their own required properties.")

table(["Schema area", "Type", "PDF destination"], [
    ("age, gender, language", "nullable scalar", "Patient profile table"),
    ("executiveSummary", "string", "Section 1"),
    ("presentingConcerns", "string[]", "Section 2"),
    ("backgroundHistory", "string", "Section 3"),
    ("clinicalThemes", "string[]", "Section 4"),
    ("clinicalObservations", "string[]", "Section 5"),
    ("riskFormulation", "object", "Section 6"),
    ("protectiveFactors", "string[]", "Section 7"),
    ("reasoningModel", "string", "Section 8"),
    ("evidenceModel", "object[]", "Section 9"),
    ("questionsAskedAndWhy", "object[]", "Section 10"),
    ("confidenceScore/rationale", "integer/string", "Section 11"),
    ("keywords", "string[]", "Section 12"),
    ("reviewerConsiderations", "string[]", "Section 13"),
    ("facilitatorSummary", "string", "Section 14"),
    ("clinicalWorkingNotes", "string", "Section 15"),
], [2800, 2100, 4460])

doc.add_heading("4.1 Risk object", level=2)
code('{\n  "level": "low | medium | high | urgent | unknown",\n  "summary": "evidence-bound narrative",\n  "currentSafetyAssessed": false,\n  "warningSigns": [],\n  "followUp": []\n}')

doc.add_heading("4.2 Why schema validation matters", level=2)
for x in [
    "Prevents missing fields from silently breaking PDF generation.",
    "Prevents unexpected properties from entering the report pipeline.",
    "Makes missing information explicit and predictable.",
    "Keeps API behavior testable across prompt and model changes.",
]: bullet(x)

page_break()
doc.add_heading("5. PDF generation", level=1)
callout("Renderer", "src/pdf.js uses PDFKit. It does not modify the original template PDF; it recreates the template's visual hierarchy and 15-section structure with dynamic pagination.")

doc.add_heading("5.1 Renderer input", level=2)
code('await createReportPdf({\n  metadata,       // form values\n  report,         // strict OpenAI JSON\n  logoPath,       // LOGO_PATH from .env\n  outputPath      // generated server path\n});')

doc.add_heading("5.2 Layout behavior", level=2)
for x in [
    "A4 portrait pages with restrained blue/green WannaTalk styling.",
    "Patient metadata appears in a fixed profile grid.",
    "Narrative sections use bordered cards; list sections use compact list rows.",
    "Long content creates additional pages automatically rather than clipping.",
    "Every page receives a consistent running header, footer, and Page X of Y label.",
    "A valid PNG or JPEG at LOGO_PATH is placed on the first page; missing/invalid logo is non-fatal.",
    "Null or blank scalars display as Not available; absent question content is explained explicitly.",
]: bullet(x)

doc.add_heading("5.3 File lifecycle", level=2)
for step in [
    "Multer writes the uploaded transcript to uploads/ using a temporary name.",
    "The generated PDF is written to output/pdf/ with a sanitized reference and random suffix.",
    "res.download returns a friendly filename to the browser.",
    "The finally block deletes the temporary uploaded transcript even if generation fails.",
    "Generated PDFs remain server-side until a retention or deletion policy is implemented.",
]: numbered(step)

callout("Production decision required", "Define and implement a generated-PDF retention policy. These reports contain sensitive personal and clinical information and should not remain indefinitely on local disk.", fill="FDEEEE", accent=RED)

page_break()
doc.add_heading("6. Configuration and startup", level=1)
doc.add_heading("6.1 Environment variables", level=2)
table(["Variable", "Required", "Example", "Purpose"], [
    ("OPENAI_API_KEY", "Yes", "sk-...", "Server-side OpenAI authentication"),
    ("OPENAI_MODEL", "No", "gpt-5-mini", "Extraction model override"),
    ("PORT", "No", "3000", "HTTP listening port"),
    ("LOGO_PATH", "No", "C:\\assets\\logo.png", "Absolute PNG/JPEG logo path"),
], [2100, 1300, 2600, 3360])

doc.add_heading("6.2 Local setup", level=2)
code('cd C:\\Jacob\\PdfExport\nnpm install\nCopy-Item .env.example .env\n# Edit .env and add OPENAI_API_KEY\nnpm start')
doc.add_paragraph("Open http://localhost:3000 and submit the form. For a PDF-only smoke test that does not call OpenAI, run npm run sample.")

doc.add_heading("6.3 Directory permissions", level=2)
for x in [
    "The Node.js process must be able to read LOGO_PATH.",
    "The process must be able to create and delete files under uploads/.",
    "The process must be able to create files under output/pdf/.",
]: bullet(x)

page_break()
doc.add_heading("7. Error handling and observability", level=1)
table(["Failure", "Current handling", "Recommended production improvement"], [
    ("Missing metadata", "400 JSON error", "Add field-level UI feedback"),
    ("Invalid/large upload", "Multer 400 error", "Log rejection reason without transcript content"),
    ("Empty transcript", "400 JSON error", "Include correlation ID"),
    ("Missing API key", "500 with configuration message", "Fail fast at server startup"),
    ("OpenAI failure", "500 and server console error", "Structured logs, retry only transient failures"),
    ("Invalid output", "JSON parsing error", "Record response ID, never raw clinical content in logs"),
    ("PDF/filesystem failure", "500 and cleanup", "Disk monitoring and alerting"),
], [2100, 3000, 4260])

doc.add_heading("7.1 Logging rules", level=2)
for x in [
    "Do log: request correlation ID, reference hash, processing duration, response ID, model name, and error class.",
    "Do not log: full transcript, patient name, contact information, extracted report JSON, or API key.",
    "Separate operational logs from audit logs and apply restricted access and retention.",
]: bullet(x)

doc.add_heading("7.2 Retry policy", level=2)
doc.add_paragraph("Retry only network timeouts, rate limits, and transient 5xx failures. Use capped exponential backoff with jitter and an idempotency/correlation strategy. Do not automatically retry validation errors, missing keys, or malformed application inputs.")

page_break()
doc.add_heading("8. Security, privacy, and clinical safeguards", level=1)
for x in [
    "Keep OPENAI_API_KEY in a server-side secret store; never expose it in HTML or browser JavaScript.",
    "Add authentication and role-based authorization before production use.",
    "Use HTTPS and avoid public access to uploads/ and output/pdf/.",
    "Apply antivirus/content scanning and stronger MIME verification for uploads.",
    "Encrypt data at rest and define transcript/PDF retention and secure deletion policies.",
    "Record consent and the lawful basis for processing health-related data.",
    "Do not log PII, transcripts, evidence excerpts, or completed reports.",
    "Require human review before the report is used for clinical decisions.",
    "Present risk as a working formulation; directly assess safety when it was not assessed in the transcript.",
    "Review applicable healthcare, privacy, professional, and cross-border data requirements with qualified counsel.",
]: bullet(x)

callout("Current scope", "The project is a functional prototype. It does not yet include authentication, database audit trails, encrypted object storage, automated retention, background job processing, or formal clinical governance.", fill="FFF4DF", accent="7A5A00")

doc.add_heading("8.1 Threat checklist", level=2)
table(["Risk", "Control"], [
    ("Unauthorized report generation", "Authentication, authorization, rate limiting"),
    ("Path manipulation", "Server-generated output path and sanitized reference"),
    ("Oversized upload", "Multer 2 MB limit"),
    ("Malicious upload", "Extension/MIME check now; add content scanning"),
    ("Secret leakage", ".env ignored; use managed secrets in production"),
    ("Sensitive disk accumulation", "Implement retention and encrypted storage"),
    ("Prompt injection in transcript", "No tools exposed; strict schema; evidence-only policy; human review"),
], [2900, 6460])

page_break()
doc.add_heading("9. Testing strategy", level=1)
doc.add_heading("9.1 Existing checks", level=2)
for x in [
    "npm run sample verifies PDF generation without API usage.",
    "Node syntax checks validate server, extraction, and PDF modules.",
    "The generated sample was rendered and visually inspected for clipping, headers, footers, and all report sections.",
    "A server smoke test confirmed the browser form returns HTTP 200.",
]: bullet(x)

doc.add_heading("9.2 Recommended automated tests", level=2)
table(["Layer", "Test cases"], [
    ("Validation", "Missing name/reference/contact/type/file; empty file; >2 MB; wrong type"),
    ("Extraction", "Missing demographics, inferred language, no questions, no safety assessment, explicit diagnosis"),
    ("Schema", "Reject extra fields, missing nested fields, invalid confidence/risk values"),
    ("PDF", "Long text pagination, absent lists, logo present/missing/invalid, special characters"),
    ("Integration", "Mock OpenAI success, timeout, 429, 5xx, empty output, invalid JSON"),
    ("Security", "Filename attacks, unauthorized requests, rate limit, sensitive-log scan"),
], [1900, 7460])

doc.add_heading("9.3 Acceptance checklist", level=2)
for x in [
    "A valid form and transcript return a readable PDF.",
    "All 15 sections appear in order.",
    "Missing information is labeled and never invented.",
    "Current safety is not claimed when unassessed.",
    "The browser never receives the OpenAI key.",
    "Temporary transcripts are removed after success and failure.",
    "Headers, footers, page numbers, tables, and long content render without clipping.",
]: bullet(x)

page_break()
doc.add_heading("10. Deployment guidance", level=1)
doc.add_heading("10.1 Minimum production topology", level=2)
for step in [
    "Deploy the Node.js service behind an HTTPS reverse proxy or managed application platform.",
    "Store OPENAI_API_KEY in the platform's secret manager.",
    "Replace local PDF storage with encrypted private object storage and time-limited downloads.",
    "Add identity, role checks, request rate limits, audit events, and correlation IDs.",
    "Set CPU, memory, request-body, timeout, and concurrency limits.",
    "Add health/readiness endpoints and centralized metrics without clinical content.",
    "Implement retention and deletion jobs for uploads and generated reports.",
]: numbered(step)

doc.add_heading("10.2 Scalability note", level=2)
doc.add_paragraph("The current endpoint performs extraction and PDF generation synchronously. For higher volume, submit a background job, store status by job ID, generate the report in a worker, and provide an authenticated one-time download when complete. This prevents long HTTP requests and enables retry, cancellation, and capacity control.")

doc.add_heading("10.3 Health checks", level=2)
table(["Check", "Should verify"], [
    ("Liveness", "Node process responds"),
    ("Readiness", "Required configuration and writable storage are available"),
    ("Dependency telemetry", "OpenAI latency/error rates without sending test clinical content"),
    ("Disk/storage", "Capacity and retention jobs are healthy"),
], [2300, 7060])

page_break()
doc.add_heading("11. Troubleshooting", level=1)
table(["Symptom", "Likely cause", "Resolution"], [
    ("OPENAI_API_KEY is not configured", ".env missing or process not restarted", "Create .env, set key, restart Node"),
    ("Report could not be generated", "OpenAI/network/schema/filesystem error", "Check server logs by error class and response ID"),
    ("400 on upload", "Missing field, empty file, type or size failure", "Check multipart names and .txt size"),
    ("Logo absent", "LOGO_PATH empty, unreadable, or invalid image", "Use an absolute readable PNG/JPEG path"),
    ("PDF has extra pages", "Long model output", "Constrain extraction length or refine PDF spacing"),
    ("npm start cannot find node", "Node.js not installed or PATH missing", "Install supported Node.js and reopen terminal"),
    ("Permission denied", "No write access to uploads/output", "Grant service account scoped directory access"),
], [2400, 3000, 3960])

doc.add_heading("12. Developer change guide", level=1)
table(["Change", "Edit"], [
    ("Add/remove extracted field", "Update src/report-schema.js, extraction policy, src/pdf.js, and tests together"),
    ("Change model", "Set OPENAI_MODEL; validate output and run extraction evals"),
    ("Change logo", "Set LOGO_PATH; no source change required"),
    ("Change PDF styling", "Edit colors, helpers, and section calls in src/pdf.js"),
    ("Change upload limit", "Edit Multer limits in src/server.js and update UI/docs"),
    ("Add authentication", "Protect POST /api/reports before Multer processes the upload"),
    ("Use database records", "Resolve metadata/transcript server-side after authorization; do not trust browser-supplied IDs alone"),
], [2800, 6560])

doc.add_heading("13. Release checklist", level=1)
for x in [
    "Run syntax, unit, integration, and PDF-render tests.",
    "Confirm the deployed model supports the configured structured output.",
    "Confirm authentication and authorization are enforced.",
    "Confirm secrets are managed outside source control.",
    "Confirm encrypted storage and retention/deletion behavior.",
    "Confirm logs contain no PII or transcript/report content.",
    "Complete clinical, privacy, and security review.",
    "Verify the final PDF visually with short, long, missing-data, and non-English transcripts.",
]: bullet(x)

doc.add_heading("14. Reference links", level=1)
for text in [
    "OpenAI API quickstart: https://platform.openai.com/docs/quickstart/make-your-first-api-request",
    "OpenAI Responses API reference: https://developers.openai.com/api/reference",
    "Project README: C:\\Jacob\\PdfExport\\README.md",
]: bullet(text)

# Keep tables from splitting individual rows and set document metadata.
for t in doc.tables:
    for row in t.rows:
        trPr = row._tr.get_or_add_trPr()
        cantSplit = OxmlElement("w:cantSplit"); trPr.append(cantSplit)

doc.core_properties.title = "WannaTalk Intake Transcript to Clinical PDF - Developer Technical Guide"
doc.core_properties.subject = "Architecture, API flow, extraction schema, PDF generation, security, testing, and deployment"
doc.core_properties.author = "WannaTalk"
doc.core_properties.keywords = "Node.js, OpenAI Responses API, PDFKit, clinical intake, technical guide"
doc.save(OUT)
print(OUT)
